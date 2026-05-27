from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
        # Add a border-like line (using underscore for simplicity in basic docx)
        p.paragraph_format.space_after = Pt(2)
        
    # --- HEADER ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = header.add_run("GOURAV KARMAKAR")
    name.bold = True
    name.font.size = Pt(24)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = (
        "Agartala, Tripura | +91 6289130024 | cinefix1.12.19@gmail.com\n"
        "LinkedIn: linkedin.com/in/gourav-karmakar-6b6a96279 | GitHub: github.com/Dhruv727876"
    )
    run = contact.add_run(contact_text)
    run.font.size = Pt(10)

    # --- SUMMARY ---
    add_section_header("Professional Summary")
    summary = doc.add_paragraph(
        "Highly motivated B.Tech Computer Science student with a strong foundation in Full Stack Web Development. "
        "Expertise in building high-performance, visually stunning digital systems using modern frameworks like Next.js and React. "
        "Proven ability to translate complex design concepts into interactive, production-ready applications. "
        "Seeking an internship to contribute technical skills in frontend and backend development."
    )
    summary.paragraph_format.space_after = Pt(12)

    # --- EDUCATION ---
    add_section_header("Education")
    edu_p = doc.add_paragraph()
    run = edu_p.add_run("Techno India University Tripura")
    run.bold = True
    edu_p.add_run("\t\t\t\t\t\tAgartala, Tripura")
    
    edu_sub = doc.add_paragraph()
    edu_sub.add_run("Bachelor of Technology in Computer Science & Engineering (Core)").italic = True
    edu_sub.add_run("\t\t\tExpected Aug 2029")
    
    doc.add_paragraph(f"CGPA: 7.52 / 10.0 (1st Semester)", style='List Bullet')
    doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # --- TECHNICAL SKILLS ---
    add_section_header("Technical Skills")
    skills = [
        ("Languages", "JavaScript (ES6+), TypeScript, HTML5, CSS3, SQL, Python"),
        ("Frontend", "React 19, Next.js 15, Vite 8, GSAP, Three.js, Tailwind CSS 4, Framer Motion, Recharts"),
        ("Backend", "Node.js, Express 5, Bun, RESTful APIs, Firebase Admin SDK, Groq SDK"),
        ("AI & Cloud", "Google Gemini 2.0 Flash (Search Grounding), NVIDIA NIM (Gemma), Groq, Firebase (Auth/Firestore), Render"),
        ("Tools", "Docker, Git, GitHub, Figma, Vercel, Postman, Axios, Zod, Multer")
    ]
    for category, items in skills:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{category}: ")
        run.bold = True
        p.add_run(items)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # --- PROJECTS ---
    add_section_header("Selected Projects")

    # Project 1: SevaSync
    p1 = doc.add_paragraph()
    run = p1.add_run("SevaSync | AI-Powered Disaster Response Platform")
    run.bold = True
    p1.add_run("\t\tReact 19, Vite 8, Groq SDK, Firebase")
    doc.add_paragraph(
        "Engineered an intelligent disaster coordination system that converts community data into actionable intelligence with high-precision volunteer matching.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Integrated Groq SDK for ultra-fast AI-powered logic and decision-making, coupled with Llama-based scene analysis for urgency scoring.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Developed a robust Node.js/Express backend integrated with Firebase Admin SDK for secure data orchestration and real-time synchronization.",
        style='List Bullet'
    )

    # Project 2: VoteSaathi
    p2 = doc.add_paragraph()
    run = p2.add_run("VoteSaathi | AI-Driven Election Assistant")
    run.bold = True
    p2.add_run("\t\t\tReact 18, Gemini 2.0 Flash, Firebase")
    doc.add_paragraph(
        "Developed a comprehensive election assistant providing real-time candidate data and voter registration guidance with Google Search grounding.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Leveraged Google Gemini 2.0 Flash and NVIDIA NIM (Gemma) to provide context-aware insights on voting procedures and policy analysis.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Implemented secure anonymous authentication and real-time data sync using Firebase, deployed on Render and Firebase Hosting.",
        style='List Bullet'
    )

    # Project 3: CORESHIFT
    p3 = doc.add_paragraph()
    run = p3.add_run("CORESHIFT | Full-Stack HR Management Platform")
    run.bold = True
    p3.add_run("\t\tNext.js 16, Bun, PostgreSQL")
    doc.add_paragraph(
        "Designed and implemented an all-in-one HRMS for scaling teams, focusing on workforce intelligence and operational speed.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Utilized Next.js for server-side rendering and Bun for high-performance runtime execution, reducing server latency significantly.",
        style='List Bullet'
    )

    # Project 4: FOLIO OS
    p4 = doc.add_paragraph()
    run = p4.add_run("FOLIO OS | Celestial Desktop Experience")
    run.bold = True
    p4.add_run("\t\t\tThree.js, GSAP 3, WebGL")
    doc.add_paragraph(
        "Engineered an immersive celestial desktop operating system experience for technical portfolios with 60fps interaction rendering.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Implemented a custom 'Aura' particle engine and physics-based card interactions using Three.js and WebGL for high visual impact.",
        style='List Bullet'
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- ACHIEVEMENTS & INTERESTS ---
    add_section_header("Achievements & Interests")
    doc.add_paragraph("Successfully developed and deployed over 20+ complex web projects featuring advanced AI integrations.", style='List Bullet')
    doc.add_paragraph("Expertise in building high-performance systems using modern tech stacks (React 19, Gemini API, Cloud Run).", style='List Bullet')
    doc.add_paragraph("Passionate about bridging the gap between AI and human-centric application design.", style='List Bullet')

    # Save the document
    file_name = "Gourav_Karmakar_Resume_Updated.docx"
    doc.save(file_name)
    print(f"Resume generated: {file_name}")

if __name__ == "__main__":
    create_resume()
